from typing import List, Iterator
from langchain.schema import Document
import docx

class SimpleDocxLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        return list(self.lazy_load())  # ✅ lazy_load()를 기본으로 사용

    def lazy_load(self) -> Iterator[Document]:  # ✅ DirectoryLoader와 호환되도록 추가
        doc = docx.Document(self.file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        yield Document(page_content=full_text, metadata={"source": self.file_path})
